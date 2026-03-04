"""
DOCX Parser using python-docx

Extracts text content from DOCX files for question bank processing.
"""

from typing import Optional
import docx
from io import BytesIO

from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXParser:
    """Parser for DOCX files"""
    
    @staticmethod
    def parse(file_content: bytes, file_name: str) -> str:
        """
        Parse DOCX file and extract text content
        
        Args:
            file_content: Raw bytes of the DOCX file
            file_name: Name of the file (for logging)
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If parsing fails
        """
        try:
            logger.info(f"Parsing DOCX file: {file_name}")
            
            # Create a BytesIO object from the file content
            docx_file = BytesIO(file_content)
            
            # Load the document
            document = docx.Document(docx_file)
            
            # Extract text from all paragraphs
            text_content = []
            
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            # Join all text with newlines
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content could be extracted from DOCX")
            
            logger.info(
                f"Successfully parsed DOCX: {file_name}",
                extra={
                    "file_name": file_name,
                    "num_paragraphs": len(document.paragraphs),
                    "num_tables": len(document.tables),
                    "text_length": len(full_text)
                }
            )
            
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX parsing error for {file_name}: {str(e)}")
            raise Exception(f"Failed to parse DOCX file: {str(e)}")
